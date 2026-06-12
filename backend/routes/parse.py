from flask import Blueprint, request, jsonify
import os, io, re, json

parse_bp = Blueprint('parse', __name__)

ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx'}
MAX_FILE_SIZE = 10 * 1024 * 1024


# ── text extraction ──────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes):
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    return '\n'.join(page.extract_text() or '' for page in reader.pages)

def extract_text_from_docx(file_bytes):
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return '\n'.join(parts)

def extract_text_from_doc(file_bytes):
    try:
        return extract_text_from_docx(file_bytes)
    except Exception:
        pass
    text_utf16 = ''
    try:
        decoded = file_bytes.decode('utf-16-le', errors='ignore')
        runs = re.findall(r'[\x20-\x7E\n\r\t]{4,}', decoded)
        text_utf16 = '\n'.join(runs)
    except Exception:
        pass
    decoded_l1 = file_bytes.decode('latin-1', errors='ignore')
    runs_l1 = re.findall(r'[\x20-\x7E]{4,}', decoded_l1)
    text_l1 = '\n'.join(runs_l1)
    text = text_utf16 if len(text_utf16) > len(text_l1) else text_l1
    text = re.sub(r'(Microsoft|Normal\.dot|Times New Roman|HYPERLINK)[^\n]*', '', text)
    return text

def extract_text(filename, file_bytes):
    ext = os.path.splitext(filename.lower())[1]
    if ext == '.pdf':  return extract_text_from_pdf(file_bytes)
    if ext == '.docx': return extract_text_from_docx(file_bytes)
    if ext == '.doc':  return extract_text_from_doc(file_bytes)
    raise ValueError(f'Unsupported: {ext}')


# ── section splitting ─────────────────────────────────────────────────────────

SECTION_PATTERNS = {
    'summary':        r'summary|objective|profile|about\s*me|career\s*objective|professional\s*summary|professional\s*profile',
    'experience':     r'work\s*experience|experience|employment|work\s*history|professional\s*experience|career\s*history',
    'education':      r'education|academic|qualification|academic\s*background|educational\s*qualification',
    'skills':         r'skills|technical\s*skills|core\s*competencies|technologies|key\s*skills|areas\s*of\s*expertise',
    'projects':       r'projects|personal\s*projects|key\s*projects|academic\s*projects|notable\s*projects',
    'certifications': r'certifications?|licen[sc]es?|courses?|achievements?|awards?\s*&?\s*certifications?|training',
    'languages':      r'languages?|language\s*proficiency|languages?\s*known',
}

def split_sections(text):
    lines = text.split('\n')
    sections = {'_header': []}
    current = '_header'
    for line in lines:
        s = line.strip()
        if s and len(s) < 55:
            matched = None
            for key, pat in SECTION_PATTERNS.items():
                if re.fullmatch(r'(?:' + pat + r')\s*:?', s, re.IGNORECASE):
                    matched = key
                    break
            if matched:
                current = matched
                sections.setdefault(current, [])
                continue
        sections.setdefault(current, []).append(line)
    return {k: '\n'.join(v).strip() for k, v in sections.items()}


# ── date helpers ──────────────────────────────────────────────────────────────

_MON  = r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
_YR   = r'\d{4}'
_DATE = rf'(?:{_MON}\.?\s*{_YR}|\d{{1,2}}/\d{{4}}|{_YR})'
_END  = rf'(?:{_DATE}|Present|Current|Now|Till\s+[Dd]ate|Ongoing|Till\s+now)'
_RANGE_RE = re.compile(rf'({_DATE})\s*[-–—to]+\s*({_END})', re.IGNORECASE)

def find_date_range(line):
    m = _RANGE_RE.search(line)
    if m:
        start = m.group(1).strip()
        end   = m.group(2).strip()
        cur   = end.lower() in ('present', 'current', 'now', 'ongoing') or 'till' in end.lower()
        return start, ('' if cur else end), cur, m
    return None, None, False, None


# ── personal info ─────────────────────────────────────────────────────────────

_TITLE_KW = [
    'developer','engineer','designer','manager','analyst','architect',
    'consultant','lead','intern','officer','executive','director',
    'scientist','researcher','specialist','coordinator','administrator',
    'programmer','technician','advisor','associate','full stack','frontend',
    'backend','software','data','product','project','ui','ux','devops',
    'cloud','mobile','web','qa','tester','sre','scrum master',
]

def extract_personal_info(text):
    lines = [l.strip() for l in text.split('\n') if l.strip()]

    email_m     = re.search(r'[\w.+-]+@[\w-]+\.[\w.]+', text)
    phone_m     = re.search(
        r'(\+?91[\s\-]?)?[6-9]\d{4}[\s\-]?\d{5}'
        r'|\+?\d{1,3}[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4}',
        text
    )
    linkedin_m  = re.search(r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+', text, re.IGNORECASE)
    # portfolio: URL that is NOT linkedin and NOT email domain
    email_domain = email_m.group(0).split('@')[1].lower() if email_m else ''
    portfolio = ''
    for pm in re.finditer(r'(?:https?://)?(?:www\.)?[\w\-]+\.(?:com|io|dev|me|co|in|net|org)(?:/[\w\-/]*)?', text, re.IGNORECASE):
        url = pm.group(0).lower()
        if 'linkedin' not in url and (not email_domain or email_domain not in url):
            portfolio = pm.group(0)
            break

    name = title = location = ''

    for line in lines[:8]:
        if (2 <= len(line.split()) <= 5
                and not re.search(r'[@\d|/\\]', line)
                and len(line) < 60
                and not re.search(r'resume|curriculum|vitae|summary|objective', line, re.IGNORECASE)):
            if not name:
                name = line.title() if line.isupper() else line
            elif not title and any(kw in line.lower() for kw in _TITLE_KW):
                title = line
                break

    if not title:
        for line in lines[1:7]:
            if any(kw in line.lower() for kw in _TITLE_KW) and len(line) < 70 and '@' not in line:
                title = line
                break

    # Location: City, State / City PIN / City, Country patterns
    loc_pats = [
        r'[A-Z][a-zA-Z\s]+,\s*[A-Z][a-zA-Z]+\s+\d{6}',   # City, State PIN
        r'[A-Z][a-zA-Z\s]+,\s*[A-Z][a-zA-Z\s]+,\s*[A-Z][a-zA-Z]+',  # City, State, Country
        r'[A-Z][a-zA-Z]+,\s*[A-Z][a-zA-Z]+',              # City, State
        r'[A-Z][a-zA-Z\s]+\d{6}',                          # City PIN
    ]
    header_text = '\n'.join(lines[:12])
    for pat in loc_pats:
        lm = re.search(pat, header_text)
        if lm:
            location = lm.group(0).strip()
            break

    return {
        'fullName':  name,
        'title':     title,
        'email':     email_m.group(0) if email_m else '',
        'phone':     phone_m.group(0).strip() if phone_m else '',
        'location':  location,
        'linkedin':  linkedin_m.group(0) if linkedin_m else '',
        'portfolio': portfolio,
    }


# ── experience parser ─────────────────────────────────────────────────────────

def parse_experience_section(text):
    if not text.strip():
        return []

    # Keep raw lines so we can detect indentation (indented = bullet/description)
    raw_lines = text.split('\n')
    entries = []
    current = None
    desc = []
    prev_clean = ''   # last non-empty stripped line (used as fallback job title)

    def flush():
        if current and (current['jobTitle'] or current['company']):
            current['description'] = '\n'.join(x for x in desc if x).strip()
            entries.append(dict(current))

    for raw in raw_lines:
        stripped = raw.strip()
        if not stripped:
            continue

        is_indented = raw.startswith((' ', '\t'))
        start, end, is_cur, m = find_date_range(stripped)

        if start:
            flush()
            # Title on same line (two-column PDF: "Senior Dev   Jan 2024 - Present")
            title_part = stripped[:m.start()].strip().strip('|•–— -') if m else ''
            # If date is on its own line, use the previous clean line as title
            if not title_part and prev_clean and not find_date_range(prev_clean)[0]:
                title_part = prev_clean
                # Remove that line from desc if it was accidentally added
                if desc and desc[-1] == prev_clean:
                    desc.pop()

            current = {
                'jobTitle':  title_part,
                'company':   '',
                'location':  '',
                'startDate': start,
                'endDate':   end,
                'current':   is_cur,
                'description': ''
            }
            desc = []

        elif current is not None:
            # Indented lines or lines starting with bullet chars = description
            is_bullet_char = bool(re.match(r'^[•\-–*▪◦\xb7\x7f]', stripped))
            clean = re.sub(r'^[•\-–*▪◦\xb7\x7f]\s*', '', stripped).strip()

            if is_indented or is_bullet_char:
                if clean:
                    desc.append(clean)
            elif len(stripped) < 100:
                if not current['company']:
                    if '|' in stripped:
                        parts = [p.strip() for p in stripped.split('|', 1)]
                        current['company']  = parts[0]
                        current['location'] = parts[1] if len(parts) > 1 else ''
                    elif not current['jobTitle']:
                        current['jobTitle'] = stripped
                    else:
                        current['company'] = stripped
                else:
                    if clean:
                        desc.append(clean)
            else:
                if clean:
                    desc.append(clean)
        else:
            # Before first date entry — might be job title line
            pass

        prev_clean = stripped

    flush()
    for i, e in enumerate(entries):
        e['id'] = i + 1
    return entries


# ── education parser ──────────────────────────────────────────────────────────

_DEG_RE = re.compile(
    r'\b(?:B\.?(?:E|Tech|Sc|Com|A|S|C|Ed|Pharm|Des|Arch|Eng)'
    r'|M\.?(?:E|Tech|Sc|Com|A|S|C|Ed|BA|CA|Eng)'
    r'|Ph\.?D|MBA|BBA|BCA|MCA|Diploma|HSC|SSC'
    r'|Bachelor(?:\s+of)?|Master(?:\s+of)?|Doctor(?:\s+of)?'
    r'|10th\s+Grade|12th\s+Grade|High\s+School|Secondary)\b',
    re.IGNORECASE
)
_GRADE_RE = re.compile(
    r'(?:CGPA|GPA|Percentage|Score)[:\s]+[\d.]+\s*%?'
    r'|[\d.]+\s*(?:CGPA|GPA|%|out\s+of\s+\d)'
    r'|[\d.]+/[\d.]+',
    re.IGNORECASE
)

def parse_education_section(text):
    if not text.strip():
        return []

    raw_lines = text.split('\n')
    entries = []
    current = None

    def flush():
        if current and (current['degree'] or current['institution']):
            entries.append(dict(current))

    for raw in raw_lines:
        line = raw.strip()
        if not line:
            continue

        deg_m   = _DEG_RE.search(line)
        grade_m = _GRADE_RE.search(line)
        start, end, is_cur, _ = find_date_range(line)
        yr_m    = re.findall(r'\b(20\d{2}|19\d{2})\b', line)

        if deg_m:
            flush()
            # Strip year/grade from degree line if present
            degree_clean = line
            if grade_m:
                degree_clean = degree_clean.replace(grade_m.group(0), '').strip().strip('|,')
            year_val = end or start or (yr_m[-1] if yr_m else '')
            if year_val:
                degree_clean = degree_clean.replace(year_val, '').strip().strip('|,-–—')
            current = {
                'degree':      degree_clean,
                'institution': '',
                'location':    '',
                'year':        year_val,
                'grade':       grade_m.group(0).strip() if grade_m else ''
            }
        elif current is not None:
            if grade_m and not current['grade']:
                current['grade'] = grade_m.group(0).strip()
            if not current['year'] and (start or yr_m):
                current['year'] = end or start or (yr_m[-1] if yr_m else '')
            if not current['institution'] and '@' not in line and len(line) < 90:
                if '|' in line:
                    parts = [p.strip() for p in line.split('|')]
                    current['institution'] = parts[0]
                    # Last part might be grade, second might be location
                    for part in parts[1:]:
                        if _GRADE_RE.search(part):
                            if not current['grade']:
                                current['grade'] = _GRADE_RE.search(part).group(0).strip()
                        elif not current['location']:
                            current['location'] = part.strip()
                elif not re.search(r'^\d', line):
                    current['institution'] = line

    flush()
    for i, e in enumerate(entries):
        e['id'] = i + 1
    return entries


# ── skills parser ─────────────────────────────────────────────────────────────

def parse_skills_section(text):
    if not text.strip():
        return []
    # Remove "Category: " prefixes
    text = re.sub(r'^[A-Za-z /]+:\s*', '', text, flags=re.MULTILINE)
    raw = re.split(r'[,\n|;:\t·•]+', text)
    skills = []
    for s in raw:
        s = s.strip().strip('–—•▪◦-*# ')
        if 1 < len(s) < 40 and s not in skills and not re.match(r'^\d+$', s):
            skills.append(s)
    return skills[:25]


# ── projects parser ───────────────────────────────────────────────────────────

_TECH_SET = {
    'react','angular','vue','svelte','node','express','python','flask','django',
    'fastapi','java','spring','typescript','javascript','html','css','tailwind',
    'bootstrap','sql','mongodb','postgresql','mysql','sqlite','redis','graphql',
    'aws','azure','gcp','docker','kubernetes','git','github','supabase','firebase',
    'flutter','swift','kotlin','c++','c#','tensorflow','pytorch','sklearn','pandas',
    'numpy','vite','webpack','razorpay','stripe','openrouter','openai','langchain',
    'next','nuxt','gatsby','remix','astro','php','laravel','ruby','rails','go',
    'rust','scala','r','bash','linux','nginx','apache',
}

def _is_tech_line(s):
    words = re.findall(r'\w+', s.lower())
    return sum(1 for w in words if w in _TECH_SET) >= 2

def parse_projects_section(text):
    if not text.strip():
        return []

    # Use raw lines to detect indentation (indented = description, not a project name)
    raw_lines = text.split('\n')
    entries = []
    current = None
    desc = []

    def flush():
        if current and current.get('projectName'):
            current['description'] = '\n'.join(x for x in desc if x).strip()
            entries.append({k: v for k, v in current.items()})

    for raw in raw_lines:
        stripped = raw.strip()
        if not stripped:
            continue

        is_indented   = raw.startswith((' ', '\t'))
        is_bullet_chr = bool(re.match(r'^[•\-–*▪◦\xb7\x7f]', stripped))
        clean = re.sub(r'^[•\-–*▪◦\xb7\x7f]\s*', '', stripped).strip()

        if is_indented or is_bullet_chr:
            # Description / bullet line
            if current:
                url_m = re.search(r'https?://\S+', clean)
                if url_m and not current['projectUrl']:
                    current['projectUrl'] = url_m.group(0)
                elif _is_tech_line(clean) and not current['techStack']:
                    current['techStack'] = clean
                elif clean:
                    desc.append(clean)
        elif len(stripped) <= 100 and not re.search(r'^https?://', stripped):
            # Non-indented, non-bullet line = new project name (possibly with tech stack after |)
            flush()
            if '|' in stripped:
                parts = stripped.split('|', 1)
                current = {
                    'projectName': parts[0].strip(),
                    'description': '',
                    'techStack':   parts[1].strip(),
                    'projectUrl':  ''
                }
            else:
                current = {
                    'projectName': stripped,
                    'description': '',
                    'techStack':   '',
                    'projectUrl':  ''
                }
            desc = []
        elif current and clean:
            desc.append(clean)

    flush()
    for i, e in enumerate(entries):
        e['id'] = i + 1
    return entries


# ── certifications parser ─────────────────────────────────────────────────────

_KNOWN_ORGS = [
    'udemy','coursera','google','aws','amazon','microsoft','oracle','cisco',
    'comptia','ibm','linkedin','edx','nptel','infosys','salesforce','adobe',
    'red hat','pmi','scrum alliance','isaca','nasscom','simplilearn','great learning',
]

def parse_certifications_section(text):
    if not text.strip():
        return []

    entries = []
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        # Strip bullet chars including DEL (0x7f used by some PDF generators)
        clean = re.sub(r'^[\x7f•\-–*▪◦\xb7]+\s*', '', line).strip()
        if len(clean) < 3:
            continue

        # Date
        date_m = re.search(rf'{_MON}\.?\s*{_YR}|{_YR}', clean, re.IGNORECASE)
        issue_date = date_m.group(0).strip() if date_m else ''

        # Issuer
        issuer = ''
        for org in _KNOWN_ORGS:
            if org in clean.lower():
                issuer = org.title()
                break

        # Fallback issuer from separators
        if not issuer:
            for sep in ('|', '—', '–', ' - '):
                if sep in clean:
                    parts = clean.split(sep)
                    candidate = parts[-1].strip().split('(')[0].strip()
                    if 2 < len(candidate) < 50:
                        issuer = candidate
                    break

        # Name = everything before issuer / date
        cert_name = clean
        if issuer and issuer.lower() in cert_name.lower():
            idx = cert_name.lower().index(issuer.lower())
            cert_name = cert_name[:idx].strip().strip('|,—–- ')
        elif '|' in cert_name:
            cert_name = cert_name.split('|')[0].strip()
        elif '—' in cert_name or '–' in cert_name:
            cert_name = re.split(r'[—–]', cert_name)[0].strip()

        # Remove date from name
        if issue_date and issue_date in cert_name:
            cert_name = cert_name.replace(issue_date, '').strip().strip('()')

        if cert_name and len(cert_name) > 2:
            entries.append({
                'id':            len(entries) + 1,
                'certName':      cert_name,
                'issuingOrg':    issuer,
                'issueDate':     issue_date,
                'credentialUrl': ''
            })

    return entries


# ── languages parser ──────────────────────────────────────────────────────────

_PROF_KW = [
    'native','fluent','proficient','intermediate','basic','beginner',
    'advanced','working','professional','mother tongue','elementary','limited',
]

def parse_languages_section(text):
    if not text.strip():
        return []

    entries = []
    for part in re.split(r'[,\n|]', text):
        part = part.strip().strip('•–—*- ')
        if not part or len(part) > 60:
            continue

        lang = prof = ''

        # "English (Fluent)" pattern
        paren_m = re.search(r'\(([^)]+)\)', part)
        if paren_m:
            prof = paren_m.group(1).strip()
            lang = part[:paren_m.start()].strip()
        else:
            for kw in _PROF_KW:
                if kw in part.lower():
                    prof = kw.title()
                    lang = re.sub(re.escape(kw), '', part, flags=re.IGNORECASE).strip().strip('|-() ')
                    break
            else:
                lang = part

        lang = lang.strip().strip('|-() ')
        if lang and 1 < len(lang) < 30:
            entries.append({
                'id':          len(entries) + 1,
                'language':    lang.title() if lang.isupper() else lang,
                'proficiency': prof.title() if prof else 'Fluent'
            })

    return entries[:8]


# ── AI extraction ─────────────────────────────────────────────────────────────

_AI_MODELS = [
    'meta-llama/llama-3.1-8b-instruct:free',
    'google/gemma-2-9b-it:free',
    'mistralai/mistral-7b-instruct:free',
    'minimax/minimax-m2.5',
]

_JSON_SHAPE = '''{
  "personalInfo":{"fullName":"","title":"","email":"","phone":"","location":"","linkedin":"","portfolio":""},
  "summary":"",
  "experience":[{"id":1,"jobTitle":"","company":"","location":"","startDate":"","endDate":"","current":false,"description":""}],
  "education":[{"id":1,"degree":"","institution":"","location":"","year":"","grade":""}],
  "skills":["skill"],
  "projects":[{"id":1,"projectName":"","description":"","techStack":"","projectUrl":""}],
  "certifications":[{"id":1,"certName":"","issuingOrg":"","issueDate":"","credentialUrl":""}],
  "languages":[{"id":1,"language":"","proficiency":"Fluent"}]
}'''

def ai_extract(text):
    import requests
    key = os.getenv('OPENROUTER_API_KEY')
    if not key:
        return None

    prompt = (
        'Extract ALL resume data into this EXACT JSON shape (no markdown, no explanation):\n'
        + _JSON_SHAPE
        + '\n\nRules:\n'
        '- Keep ALL experience and education entries (never skip)\n'
        '- ids: sequential integers starting at 1\n'
        '- skills: max 25, individual skill names only\n'
        '- experience description: bullet points joined by newlines\n'
        '- Empty string "" for any missing field. Never invent data.\n\n'
        'RESUME TEXT:\n' + text[:5500]
    )

    for model in _AI_MODELS:
        try:
            resp = requests.post(
                'https://openrouter.ai/api/v1/chat/completions',
                headers={
                    'Authorization': f'Bearer {key}',
                    'Content-Type': 'application/json',
                    'HTTP-Referer': 'https://resumate-ai.com',
                    'X-Title': 'ResuMate AI'
                },
                json={
                    'model': model,
                    'messages': [
                        {'role': 'system', 'content': 'You are a resume parser. Return ONLY valid JSON, no markdown, no explanations.'},
                        {'role': 'user',   'content': prompt}
                    ],
                    'max_tokens': 2800,
                    'temperature': 0.1
                },
                timeout=40
            )
            if resp.status_code != 200:
                print(f'[PARSE] {model} HTTP {resp.status_code}')
                continue
            raw = resp.json()['choices'][0]['message']['content'].strip()
            raw = raw.replace('```json', '').replace('```', '').strip()
            s, e = raw.find('{'), raw.rfind('}')
            if s < 0 or e <= s:
                continue
            result = json.loads(raw[s:e+1])
            print(f'[PARSE] AI ok: {model}')
            return result
        except Exception as ex:
            print(f'[PARSE] {model} error: {ex}')
            continue

    return None


# ── normalize to editor shape ─────────────────────────────────────────────────

def normalize(data):
    out = {
        'personalInfo': {
            'fullName':'','title':'','email':'','phone':'',
            'location':'','linkedin':'','portfolio':''
        },
        'summary':'','experience':[],'education':[],
        'skills':[],'projects':[],'certifications':[],'languages':[]
    }

    pi = data.get('personalInfo') or {}
    for k in out['personalInfo']:
        out['personalInfo'][k] = str(pi.get(k) or '')

    out['summary'] = str(data.get('summary') or '')
    out['skills']  = [str(s) for s in (data.get('skills') or []) if str(s).strip()][:25]

    list_shapes = {
        'experience':     {'jobTitle':'','company':'','location':'','startDate':'','endDate':'','current':False,'description':''},
        'education':      {'degree':'','institution':'','location':'','year':'','grade':''},
        'projects':       {'projectName':'','description':'','techStack':'','projectUrl':''},
        'certifications': {'certName':'','issuingOrg':'','issueDate':'','credentialUrl':''},
        'languages':      {'language':'','proficiency':'Fluent'},
    }
    for sec, shape in list_shapes.items():
        result = []
        for i, item in enumerate(data.get(sec) or []):
            if not isinstance(item, dict):
                continue
            entry = {'id': i + 1}
            for field, default in shape.items():
                val = item.get(field, default)
                entry[field] = bool(val) if isinstance(default, bool) else str(val or '')
            result.append(entry)
        out[sec] = result

    return out


# ── POST /api/parse/resume ────────────────────────────────────────────────────

@parse_bp.route('/resume', methods=['POST', 'OPTIONS'])
def parse_resume():
    if request.method == 'OPTIONS':
        return jsonify({'status': 'ok'}), 200

    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'}), 400

        file     = request.files['file']
        filename = file.filename or ''
        ext      = os.path.splitext(filename.lower())[1]

        if ext not in ALLOWED_EXTENSIONS:
            return jsonify({'success': False,
                            'error': f'Unsupported type "{ext}". Upload PDF, DOC, or DOCX.'}), 400

        file_bytes = file.read()
        if len(file_bytes) > MAX_FILE_SIZE:
            return jsonify({'success': False, 'error': 'File too large (max 10 MB)'}), 400
        if not file_bytes:
            return jsonify({'success': False, 'error': 'File is empty'}), 400

        try:
            text = extract_text(filename, file_bytes)
        except Exception as ex:
            print(f'[PARSE] text extract failed: {ex}')
            return jsonify({'success': False,
                            'error': 'Could not read this file. May be corrupted or a scanned image.'}), 422

        text = re.sub(r'\n{3,}', '\n\n', text).strip()
        if len(text) < 50:
            return jsonify({'success': False,
                            'error': 'No readable text found. Use a text-based (non-scanned) file.'}), 422

        print(f'[PARSE] {filename}: {len(text)} chars')

        # 1. Try AI
        ai_used = True
        data = ai_extract(text)

        # 2. Full rule-based fallback
        if not data:
            ai_used = False
            sections = split_sections(text)
            data = {
                'personalInfo':   extract_personal_info(text),
                'summary':        sections.get('summary', '')[:800],
                'experience':     parse_experience_section(sections.get('experience', '')),
                'education':      parse_education_section(sections.get('education', '')),
                'skills':         parse_skills_section(sections.get('skills', '')),
                'projects':       parse_projects_section(sections.get('projects', '')),
                'certifications': parse_certifications_section(sections.get('certifications', '')),
                'languages':      parse_languages_section(sections.get('languages', '')),
            }

        resume_data = normalize(data)

        pi = resume_data['personalInfo']
        print(
            f'[PARSE] done ai={ai_used} '
            f'name="{pi["fullName"]}" title="{pi["title"]}" loc="{pi["location"]}" '
            f'exp={len(resume_data["experience"])} edu={len(resume_data["education"])} '
            f'skills={len(resume_data["skills"])} proj={len(resume_data["projects"])} '
            f'cert={len(resume_data["certifications"])} lang={len(resume_data["languages"])}'
        )

        return jsonify({
            'success': True,
            'resumeData': resume_data,
            'aiUsed': ai_used,
            'textLength': len(text)
        })

    except Exception as ex:
        import traceback; traceback.print_exc()
        return jsonify({'success': False, 'error': str(ex)}), 500
